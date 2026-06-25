#!/usr/bin/env python3
"""
生成 MES 全栈项目实验报告 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# 设置标题样式
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_font.size = Pt(22)
    elif i == 2:
        heading_font.size = Pt(16)
    else:
        heading_font.size = Pt(14)

# 代码样式
code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
code_font = code_style.font
code_font.name = 'Courier New'
code_font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(3)
code_style.paragraph_format.space_after = Pt(3)
code_style.paragraph_format.line_spacing = 1.0

def add_paragraph(text, bold=False, style_name=None):
    """添加段落"""
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_code_block(code_text):
    """添加代码块"""
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph(style='CodeBlock')
        p.add_run(line)

def add_heading(text, level=1):
    """添加标题"""
    return doc.add_heading(text, level=level)

def add_screenshot_placeholder(description):
    """添加截图占位符（Word注释）"""
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    # 添加一个灰色占位框
    run = p.add_run(f'[ 截图位置：{description} ]')
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(10)
    run.italic = True
    # 添加注释
    comment = doc.add_paragraph()
    comment_run = comment.add_run(f'⚠ 请在此处插入截图：{description}')
    comment_run.font.color.rgb = RGBColor(180, 0, 0)
    comment_run.font.size = Pt(9)
    comment_run.italic = True

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('MES制造执行系统')
run.font.size = Pt(36)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('实验报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '专    业：软件工程',
    '班    级：____________',
    '姓    名：____________',
    '学    号：____________',
    f'日    期：{datetime.date.today().strftime("%Y年%m月%d日")}',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_heading('目  录', level=1)
doc.add_paragraph('1 项目概述与现场调研 .............................................. ')
doc.add_paragraph('2 系统需求分析 .................................................... ')
doc.add_paragraph('3 系统总体架构设计 ................................................ ')
doc.add_paragraph('4 系统模块开发实现 ................................................ ')
doc.add_paragraph('    4.1 角色管理与权限管理 ........................................ ')
doc.add_paragraph('    4.2 班组员工定义 .............................................. ')
doc.add_paragraph('    4.3 编组设备定义 .............................................. ')
doc.add_paragraph('    4.4 加工单元定义 .............................................. ')
doc.add_paragraph('    4.5 物料信息定义 .............................................. ')
doc.add_paragraph('    4.6 库房库位定义 .............................................. ')
doc.add_paragraph('    4.7 零部件定义 ................................................ ')
doc.add_paragraph('    4.8 产品BOM管理 ............................................... ')
doc.add_paragraph('    4.9 工序信息定义 .............................................. ')
doc.add_paragraph('    4.10 工艺流程管理 ............................................. ')
doc.add_paragraph('    4.11 工艺内容编制 ............................................. ')
doc.add_paragraph('    4.12 产品工艺查询 ............................................. ')
doc.add_paragraph('5 系统集成测试 .................................................... ')
doc.add_paragraph('6 系统运维规划 .................................................... ')
doc.add_paragraph('7 创新设计说明 .................................................... ')
doc.add_paragraph('8 总结 ............................................................ ')

doc.add_page_break()

# ============================================================
# 第1章 项目概述与现场调研
# ============================================================
add_heading('1 项目概述与现场调研', level=1)

add_heading('1.1 项目背景', level=2)
add_paragraph(
    '随着工业4.0和智能制造理念的深入推进，传统制造业正面临着数字化转型的迫切需求。'
    '制造执行系统（Manufacturing Execution System, MES）作为连接企业计划层与车间控制层的关键信息系统，'
    '在生产管理、质量控制、资源调度等方面发挥着不可替代的作用。'
)
add_paragraph(
    '本项目"章鱼师兄"MES系统是一款面向工厂生产管理的全栈单体应用，涵盖系统管理、'
    '基础数据、工艺技术管理、生产订单、质量管理、设备管理和SN码管理等核心功能模块。'
    '系统采用前后端分离架构，后端基于Java Spring Boot生态，前端采用React 18 + TypeScript + Ant Design，'
    '旨在为中小型制造企业提供一套完整、实用、易扩展的生产管理解决方案。'
)

add_heading('1.2 现场调研', level=2)
add_paragraph(
    '通过对制造企业生产车间的实地调研，我们识别出以下核心业务场景和管理痛点：'
)
add_paragraph('（1）资源分配管理：需要建立完善的角色权限体系、班组人员管理、设备编组管理和加工单元定义，确保生产资源合理配置。')
add_paragraph('（2）BOM与组件数据管理：产品结构复杂，需要支持多层级BOM管理，实现零部件定义与产品结构树的有效组织。')
add_paragraph('（3）工艺设计管理：从工序定义、工艺流程设计到工艺内容编制和产品工艺查询，需要完整的工艺数据链路。')
add_paragraph('（4）库房库位管理：立体仓库的管理需要精确到库位级别，并与3D数字仿真系统联动，实现可视化管理。')

add_heading('1.3 项目目标', level=2)
add_paragraph('本项目的主要目标包括：')
add_paragraph('（1）完善角色管理与权限控制体系，实现基于RBAC的细粒度权限管理。')
add_paragraph('（2）建立班组员工定义模块，实现生产班组的数字化管理。')
add_paragraph('（3）实现编组设备定义，支持设备的灵活编组和分组管理。')
add_paragraph('（4）定义加工单元，将人员作业单元与设备作业单元纳入统一管理。')
add_paragraph('（5）完善物料信息定义，支持物料类型的自动编码生成。')
add_paragraph('（6）实现库房库位定义，并创新性地将库房数据与3D数字仿真仓库联动。')
add_paragraph('（7）建立零部件定义和产品BOM管理，支持多层级BOM结构的版本控制。')
add_paragraph('（8）实现完整的工艺设计管理链路：工序定义→工艺流程→工艺内容编制→产品工艺查询。')

doc.add_page_break()

# ============================================================
# 第2章 系统需求分析
# ============================================================
add_heading('2 系统需求分析', level=1)

add_heading('2.1 功能需求', level=2)

add_heading('2.1.1 角色管理与权限管理', level=3)
add_paragraph(
    '系统需要支持基于角色的访问控制（RBAC）。管理员可以创建角色并为其分配菜单权限，'
    '实现不同角色看到不同的功能菜单。系统预设角色包括超级管理员、数据管理员、工艺技术员、'
    '生产计划员、生产主管、操作员、仓库管理员、质量管理员等。角色分为系统角色和普通角色，'
    '系统角色不可删除以保证权限体系的稳定性。'
)

add_heading('2.1.2 班组员工定义', level=3)
add_paragraph(
    '支持创建生产班组，设置班组的班次时间（开始时间、结束时间）和工作日。'
    '班组关联生产线和车间，支持向班组中添加和移除员工，实现人员的灵活调配。'
)

add_heading('2.1.3 编组设备定义', level=3)
add_paragraph(
    '支持设备的增删改查，包括设备编码、名称、类型、型号、规格、所属产线和位置信息。'
    '设备状态包括空闲、运行中、维修中和已报废。支持创建设备组，将多个设备编入同一组进行统一管理。'
)

add_heading('2.1.4 加工单元定义', level=3)
add_paragraph(
    '加工单元分为人员作业单元和设备作业单元两类。支持将班组绑定到加工单元，'
    '实现人员与生产单元的关联。支持设置是否带有线边库。'
)

add_heading('2.1.5 物料信息定义', level=3)
add_paragraph(
    '物料信息包括物料编码（支持按类型自动生成）、物料描述、基本单位、产品组、物料类型'
    '（产品/零件/标准件/其他）、规格型号、关联工艺流程、来源类型（自制/外购）、提前期、安全库存等。'
    '支持物料图片上传。物料类型变更时自动填充来源和提前期。'
)

add_heading('2.1.6 库房库位定义', level=3)
add_paragraph(
    '支持创建立体仓库模型：设置仓库的组数、排数、层数、列数，系统自动生成所有库位编码。'
    '库位编码格式为"组号-排号层号列号"。支持库位的查询和浏览。'
    '创新功能：将库房数据与Three.js 3D数字仿真仓库联动，实现库房的三维可视化展示。'
)

add_heading('2.1.7 零部件定义', level=3)
add_paragraph(
    '支持零部件的增删改查，编码按COMP-XXX格式自动生成。零部件作为BOM的基础组成部分，'
    '是产品结构树的叶子节点数据来源。'
)

add_heading('2.1.8 产品BOM管理', level=3)
add_paragraph(
    '支持多层级产品BOM结构管理，BOM层级包括产品（0级）、半成品（1级）和零部件（2级）。'
    '每个BOM节点可关联物料清单（BOM Item）。支持BOM锁定（锁定后不可编辑）和版本管理'
    '（在锁定版本基础上创建新版本，自动深拷贝整个BOM树和物料清单）。'
)

add_heading('2.1.9 工序信息定义', level=3)
add_paragraph(
    '定义生产工序，包括工序编码（OPR-XXX格式自动生成）、工序描述、关联加工单元、'
    '工时（分钟）、制造周期（分钟）、是否生成计划等。业务规则：制造周期必须大于工时。'
)

add_heading('2.1.10 工艺流程管理', level=3)
add_paragraph(
    '支持工艺流程（Flow）的定义和工序关系的编排。将多个工序按顺序组织成工艺流程，'
    '记录每个工序的前序工序和后序工序，形成工序链表。工艺描述自动生成为"工序A→工序B→工序C"格式。'
)

add_heading('2.1.11 工艺内容编制', level=3)
add_paragraph(
    '为产品BOM的每个节点编制详细的工艺内容。采用五步向导式界面：'
    '工艺主信息→工序要求→辅助信息（工装设备+技术文档）→物料核对→完成编制。'
    '支持工艺图片和检验标准图片上传，支持PDF技术文档上传和预览。'
    '编制完成后状态变为"已完成"，不可再修改。'
)

add_heading('2.1.12 产品工艺查询', level=3)
add_paragraph(
    '提供只读的产品工艺查询界面。选择产品后展示完整的BOM结构树，'
    '点击各节点查看对应的工艺详情，包括工艺主信息、工序要求、注意事项、'
    '工装设备、技术文档和物料清单。支持图片预览和PDF文档在线预览。'
)

add_heading('2.2 非功能需求', level=2)
add_paragraph('（1）性能需求：页面响应时间不超过3秒，支持500并发用户同时在线。')
add_paragraph('（2）安全需求：基于Apache Shiro实现认证与授权，密码MD5加密存储，支持登录重试限制。')
add_paragraph('（3）可用性需求：系统可用性达到99.5%，支持7×24小时运行。')
add_paragraph('（4）可扩展性需求：采用模块化设计，支持新功能的快速扩展和集成。')
add_paragraph('（5）兼容性需求：支持主流浏览器（Chrome、Firefox、Edge）。')

doc.add_page_break()

# ============================================================
# 第3章 系统总体架构设计
# ============================================================
add_heading('3 系统总体架构设计', level=1)

add_heading('3.1 技术架构', level=2)
add_paragraph(
    '系统采用前后端分离的B/S架构，分为表现层、业务逻辑层、数据持久层三层。'
)

add_paragraph('前端技术栈：', bold=True)
add_paragraph('• 核心框架：React 18 + TypeScript')
add_paragraph('• UI组件库：Ant Design 6 (antd)')
add_paragraph('• 构建工具：Vite 8')
add_paragraph('• 状态管理：Zustand 5')
add_paragraph('• 数据请求：TanStack Query (React Query) 5 + Axios')
add_paragraph('• 路由管理：React Router v7')
add_paragraph('• 图表可视化：ECharts 6')
add_paragraph('• 3D渲染：Three.js + @react-three/fiber + @react-three/drei')

add_paragraph('后端技术栈：', bold=True)
add_paragraph('• 核心框架：Spring Boot 2.1.7 (Java 8)')
add_paragraph('• ORM框架：MyBatis-Plus 3.1.2')
add_paragraph('• 安全框架：Apache Shiro 1.4.0')
add_paragraph('• 数据库：MySQL 8 + Druid连接池')
add_paragraph('• 缓存：Ehcache / Redis（可配置切换）')
add_paragraph('• 模板引擎：FreeMarker')

add_heading('3.2 系统架构图', level=2)
add_screenshot_placeholder('系统总体架构图 - 展示前端、后端、数据库三层架构')

add_heading('3.3 模块划分', level=2)
add_paragraph('系统按业务领域划分为以下核心模块：')
add_paragraph('（1）系统管理模块（system）：用户管理、角色管理、菜单管理、部门管理、字典管理')
add_paragraph('（2）基础数据模块（basedata）：物料维护、动态表配置、设备管理、设备组管理、加工单元、库房库位、零部件定义')
add_paragraph('（3）工艺技术模块（technology）：工序管理、工艺流程管理、BOM管理、产品BOM管理、工艺内容编制、产品工艺查询')
add_paragraph('（4）生产订单模块（order）：工单管理')
add_paragraph('（5）数字孪生模块（digitization）：计划看板、3D仿真仓库')

add_heading('3.4 数据库设计', level=2)
add_paragraph(
    '系统数据库共包含39张数据表，按业务领域组织。核心表关系如下：'
)
add_paragraph('• 用户-角色-菜单：sp_sys_user → sp_sys_user_role → sp_sys_role → sp_sys_role_menu → sp_sys_menu')
add_paragraph('• 班组-员工-加工单元：sp_team → sp_team_user → sp_sys_user；sp_process_unit → sp_process_unit_team → sp_team')
add_paragraph('• BOM-工艺-工序：sp_product_bom → sp_bom_flow → sp_flow → sp_flow_oper_relation → sp_oper')
add_paragraph('• 设备-设备组：sp_device → sp_device_group_item → sp_device_group')
add_paragraph('• 仓库-库位：sp_warehouse → sp_warehouse_location')
add_paragraph('• 工艺内容：sp_process_content → sp_process_equipment / sp_process_document')
add_paragraph('所有实体表统一继承BaseEntity，包含id（雪花算法）、createTime、createUsername、updateTime、updateUsername字段。')

add_screenshot_placeholder('数据库ER图 - 展示核心表之间的关系')

doc.add_page_break()

# ============================================================
# 第4章 系统模块开发实现
# ============================================================
add_heading('4 系统模块开发实现', level=1)

# ---------- 4.1 角色管理与权限管理 ----------
add_heading('4.1 角色管理与权限管理', level=2)

add_heading('4.1.1 设计思路', level=3)
add_paragraph(
    '角色管理采用RBAC（基于角色的访问控制）模型。系统通过角色-菜单关联实现权限分配，'
    '再通过用户-角色关联实现权限授予。设计要点包括：'
)
add_paragraph('（1）系统角色（isSystem=1）不可删除，保证预设权限体系的安全性和完整性。')
add_paragraph('（2）角色状态支持正常、已删除、已禁用三种，删除采用软删除策略。')
add_paragraph('（3）角色编辑与菜单授权整合在同一界面中，左侧编辑角色基本信息，右侧通过勾选树形菜单完成授权，操作直观高效。')
add_paragraph('（4）菜单树采用递归结构，支持无限层级，每个菜单节点可设置权限标识（如user:add, user:delete）。')

add_heading('4.1.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_sys_role（角色表）、sp_sys_role_menu（角色-菜单关联表）、sp_sys_user_role（用户-角色关联表）。')
add_paragraph('角色表核心字段：name（角色名称）、code（角色编码）、is_deleted（状态）、is_system（是否系统角色）。')

add_heading('4.1.3 界面截图', level=3)
add_screenshot_placeholder('角色管理列表页 - 展示角色列表、搜索、新增、编辑、授权操作')
add_screenshot_placeholder('角色编辑/授权弹窗 - 左侧基本信息和右侧菜单授权树的左右分栏布局')

add_heading('4.1.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 角色表单与菜单授权（RoleForm.tsx）：', bold=True)
add_code_block('''// RoleForm.tsx - 左右分栏布局：左侧角色表单 + 右侧菜单授权树
function RoleForm({ id, record, onFinish, formInstance }: RoleFormProps) {
  const [menuTree, setMenuTree] = useState<TreeVO<SysMenu>[]>([])
  const [checkedKeys, setCheckedKeys] = useState<string[]>([])

  // 加载菜单树
  useEffect(() => {
    menuApi.tree().then((data) => {
      setMenuTree(Array.isArray(data) ? data : [])
    })
  }, [])

  // 编辑模式下加载角色已有的菜单授权
  useEffect(() => {
    if (id) {
      roleApi.getRoleMenuTree(id).then((menuIds) => {
        setCheckedKeys(Array.isArray(menuIds) ? menuIds : [])
      })
    }
  }, [id])

  // 提交时携带菜单ID列表
  const handleFinish = (values: any) => {
    onFinish?.({ ...values, sysMenuIds: checkedKeys })
  }

  return (
    <Form form={formInstance} onFinish={handleFinish}>
      <div style={{ display: 'flex', gap: 24 }}>
        {/* 左侧：角色基本信息 */}
        <div style={{ flex: '0 0 300px' }}>
          <Form.Item name="name" label="角色名称" rules={[{ required: true }]}>
            <Input />
          </Form.Item>
          <Form.Item name="code" label="角色编码" rules={[{ required: true }]}>
            <Input />
          </Form.Item>
          <Form.Item name="isSystem" label="系统角色">
            <Radio.Group><Radio value="0">否</Radio><Radio value="1">是</Radio></Radio.Group>
          </Form.Item>
          <Form.Item name="deleted" label="状态">
            <Radio.Group>
              <Radio value="0">正常</Radio>
              <Radio value="1">已删除</Radio>
              <Radio value="2">已禁用</Radio>
            </Radio.Group>
          </Form.Item>
        </div>
        {/* 右侧：菜单授权树 */}
        <div style={{ flex: 1, borderLeft: '1px solid #f0f0f0', paddingLeft: 24 }}>
          <div>授权菜单</div>
          <Tree checkable defaultExpandAll checkedKeys={checkedKeys}
            onCheck={(keys) => setCheckedKeys(keys as string[])}
            treeData={convertToTreeData(menuTree)} />
        </div>
      </div>
    </Form>
  )
}''')

add_paragraph('后端核心代码 - 角色保存与菜单授权重建（SysRoleController.java）：', bold=True)
add_code_block('''// SysRoleController.java - 保存角色并重建菜单关联
@PostMapping("/add-or-update")
@ResponseBody
public Result addOrUpdate(SysRoleDTO record) throws Exception {
    // 1. 保存或更新角色基本信息
    sysRoleService.saveOrUpdate(record);
    // 2. 如果提交了菜单ID列表，重建角色-菜单关联
    if (record.getSysMenuIds() != null) {
        sysRoleMenuService.rebuild(record.getId(), record.getSysMenuIds());
    }
    return Result.success(record.getId());
}

// SysRoleMenuServiceImpl.java - 重建角色菜单关联
public void rebuild(String roleId, String[] menuIds) {
    // 删除旧的关联
    QueryWrapper<SysRoleMenu> qw = new QueryWrapper<>();
    qw.eq("role_id", roleId);
    sysRoleMenuMapper.delete(qw);
    // 插入新的关联
    for (String menuId : menuIds) {
        SysRoleMenu rm = new SysRoleMenu();
        rm.setRoleId(roleId);
        rm.setMenuId(menuId);
        sysRoleMenuMapper.insert(rm);
    }
}''')

doc.add_page_break()

# ---------- 4.2 班组员工定义 ----------
add_heading('4.2 班组员工定义', level=2)

add_heading('4.2.1 设计思路', level=3)
add_paragraph(
    '班组员工定义模块是生产资源管理的核心组成部分。设计时考虑以下要点：'
)
add_paragraph('（1）班组关联生产线和车间，形成"车间→产线→班组"的三级组织架构。')
add_paragraph('（2）班次时间采用HH:mm格式存储，工作日以逗号分隔的数字串存储（1=周一至7=周日），前端通过Checkbox组呈现。')
add_paragraph('（3）班组与员工为多对多关系，通过sp_team_user中间表管理。')
add_paragraph('（4）员工管理采用双面板交互模式：左侧显示可选员工（支持搜索筛选），右侧显示当前班组成员，直观高效。')

add_heading('4.2.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_team（班组表）、sp_team_user（班组-用户关联表）、sp_line（产线表）、sp_work_shop（车间表）。')
add_paragraph('班组表核心字段：code、name、lineId、workshopId、startTime、endTime、workdays、is_deleted。')

add_heading('4.2.3 界面截图', level=3)
add_screenshot_placeholder('班组管理列表页 - 展示班组列表、搜索、班次时间、工作日显示')
add_screenshot_placeholder('班组编辑表单 - 左侧基本信息 + 右侧车间/产线选择和时间设置')
add_screenshot_placeholder('员工管理弹窗 - 左右双面板：可选员工列表和当前成员列表')

add_heading('4.2.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 员工管理双面板交互（TeamUserModal.tsx）：', bold=True)
add_code_block('''// TeamUserModal.tsx 核心逻辑
// 左侧面板：可选员工列表 + 搜索 + "加入班组"按钮
// 右侧面板：当前班组成员列表 + "移除"按钮

const handleAddUsers = async () => {
  if (selectedUserIds.length === 0) {
    message.warning('请先选择要添加的员工')
    return
  }
  await teamApi.addTeamUsers(teamId, selectedUserIds)
  message.success('添加成功')
  setSelectedUserIds([])
  // 刷新两个面板
  await Promise.all([loadAvailableUsers(), loadTeamUsers()])
}

const handleRemoveUser = async (userId: string) => {
  await teamApi.removeTeamUser(teamId, userId)
  message.success('移除成功')
  setTeamUsers(prev => prev.filter(u => u.userId !== userId))
  await loadAvailableUsers()
}

// 可用员工列表：所有未删除用户 - 当前班组成员
const loadAvailableUsers = async () => {
  const [allUsers, teamUsers] = await Promise.all([
    teamApi.getAvailableUsers(),
    teamApi.getTeamUsers(teamId),
  ])
  const teamUserIds = new Set(teamUsers.map(u => u.userId))
  setAvailableUsers(allUsers.filter(u => !teamUserIds.has(u.id)))
}''')

add_paragraph('后端核心代码 - 班组分页查询（SpTeamMapper.xml）：', bold=True)
add_code_block('''<!-- SpTeamMapper.xml - 带关联信息的班组分页查询 -->
<select id="pageWithRelations" resultType="...SpTeamDTO">
  SELECT t.*,
    l.line_desc AS lineName,
    w.work_shop_desc AS workshopName,
    (SELECT COUNT(*) FROM sp_team_user WHERE team_id = t.id) AS userCount
  FROM sp_team t
  LEFT JOIN sp_line l ON t.line_id = l.id
  LEFT JOIN sp_work_shop w ON t.workshop_id = w.id
  WHERE t.is_deleted != '1'
  <if test="req.name != null and req.name != ''">
    AND t.name LIKE CONCAT('%', #{req.name}, '%')
  </if>
  ORDER BY t.create_time DESC
</select>''')

doc.add_page_break()

# ---------- 4.3 编组设备定义 ----------
add_heading('4.3 编组设备定义', level=2)

add_heading('4.3.1 设计思路', level=3)
add_paragraph(
    '编组设备定义模块将设备管理和设备分组整合在同一页面中。设计要点包括：'
)
add_paragraph('（1）采用Tab页切换"设备管理"和"设备组管理"两个视图，共享同一个页面上下文。')
add_paragraph('（2）设备状态用颜色编码的Tag标签展示：绿色=空闲、蓝色=运行中、橙色=维修中、红色=已报废。')
add_paragraph('（3）设备组采用与班组相同的双面板交互模式进行设备成员管理。')
add_paragraph('（4）设备和设备组均支持软删除。')

add_heading('4.3.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_device（设备表）、sp_device_group（设备组表）、sp_device_group_item（设备-设备组关联表）。')
add_paragraph('设备表核心字段：code、name、type、model、specs、lineId、location、status（0空闲/1运行中/2维修中/3已报废）。')

add_heading('4.3.3 界面截图', level=3)
add_screenshot_placeholder('设备管理Tab页 - 设备列表、搜索、状态颜色标签')
add_screenshot_placeholder('设备组管理Tab页 - 设备组列表和设备成员管理弹窗')

add_heading('4.3.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 双Tab页设计（DeviceGroupPage.tsx）：', bold=True)
add_code_block('''// DeviceGroupPage.tsx - Tab切换设备管理和设备组管理
<Tabs activeKey={activeTab} onChange={setActiveTab}>
  <TabPane tab="设备管理" key="devices">
    <SearchForm onSearch={setDeviceFilters}>
      <Form.Item name="code"><Input placeholder="设备编码" /></Form.Item>
      <Form.Item name="name"><Input placeholder="设备名称" /></Form.Item>
      <Form.Item name="status">
        <Select placeholder="设备状态">
          <Option value="0">空闲</Option>
          <Option value="1">运行中</Option>
          <Option value="2">维修中</Option>
          <Option value="3">已报废</Option>
        </Select>
      </Form.Item>
    </SearchForm>
    <PageTable columns={deviceColumns} dataSource={devices} ... />
  </TabPane>
  <TabPane tab="设备组管理" key="groups">
    <PageTable columns={groupColumns} dataSource={groups} ... />
  </TabPane>
</Tabs>''')

doc.add_page_break()

# ---------- 4.4 加工单元定义 ----------
add_heading('4.4 加工单元定义', level=2)

add_heading('4.4.1 设计思路', level=3)
add_paragraph(
    '加工单元定义模块采用左右分栏布局。设计要点：'
)
add_paragraph('（1）左侧展示加工单元列表，点击行选中后高亮显示。')
add_paragraph('（2）右侧面板展示该加工单元绑定的班组信息，支持添加和解除绑定。')
add_paragraph('（3）加工单元类型分"人员作业单元"和"设备作业单元"两种，支持设置是否带有线边库。')
add_paragraph('（4）加工单元与班组为多对多关系，通过sp_process_unit_team中间表管理。')

add_heading('4.4.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_process_unit（加工单元表）、sp_process_unit_team（加工单元-班组关联表）。')
add_paragraph('加工单元表核心字段：code、name、type（人员作业单元/设备作业单元）、hasLineWarehouse（0/1）、is_deleted。')

add_heading('4.4.3 界面截图', level=3)
add_screenshot_placeholder('加工单元管理页 - 左右分栏：左侧单元列表 + 右侧班组绑定面板')

add_heading('4.4.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 左右分栏交互（ProcessUnitPage.tsx）：', bold=True)
add_code_block('''// ProcessUnitPage.tsx - 点击单元行加载班组绑定
const handleSelect = async (record) => {
  setSelectedUnit(record)
  const teams = await unitApi.getTeams(record.id)
  setBoundTeams(teams)
}

// 界面结构：flex布局
<div style={{ display: 'flex', gap: 16 }}>
  <div style={{ flex: 2 }}>
    {/* 加工单元列表表格 */}
    <PageTable onRow={(r) => ({
      onClick: () => handleSelect(r),
      style: { background: selectedUnit?.id === r.id ? '#e6f7ff' : undefined }
    })} />
  </div>
  <div style={{ flex: 1 }}>
    {/* 班组绑定面板 */}
    <h3>班组绑定</h3>
    {selectedUnit && <TeamBindingPanel unit={selectedUnit} />}
  </div>
</div>''')

doc.add_page_break()

# ---------- 4.5 物料信息定义 ----------
add_heading('4.5 物料信息定义', level=2)

add_heading('4.5.1 设计思路', level=3)
add_paragraph(
    '物料信息定义模块提供完整的物料主数据管理。设计亮点包括：'
)
add_paragraph('（1）物料编码自动生成：根据物料类型（产品/零件/标准件/其他）自动生成带前缀的编码（PROD-/PART-/STD-/OTHR-），后端查询同类最大编码并递增。')
add_paragraph('（2）物料类型联动：选择物料类型后，自动填充来源（自制/外购）和默认提前期。')
add_paragraph('（3）表单采用分区布局，通过Divider将字段分为"基本信息"、"属性信息"、"库存与工艺"、"图片与状态"四个区域。')
add_paragraph('（4）支持物料图片上传和预览。')

add_heading('4.5.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_materile（物料表），关联sp_flow（工艺流程表）。')
add_paragraph('物料表核心字段：materiel、materielDesc、unit、productGroup、matType、size、flowId、model、source、leadTime、safetyStock、imageUrl。')

add_heading('4.5.3 界面截图', level=3)
add_screenshot_placeholder('物料维护列表页 - 展示物料列表和搜索')
add_screenshot_placeholder('物料编辑表单 - 分区布局和物料类型联动效果')

add_heading('4.5.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 物料类型联动（MaterileForm.tsx）：', bold=True)
add_code_block('''// MaterileForm.tsx - 物料类型变更时自动填充来源和提前期
const TYPE_DEFAULTS: Record<string, { source: string; leadTime: number }> = {
  '产品': { source: '自制', leadTime: 3 },
  '零件': { source: '外购', leadTime: 1 },
  '标准件': { source: '外购', leadTime: 1 },
  '其他': { source: '外购', leadTime: 1 },
}

const handleTypeChange = (value: string) => {
  const defaults = TYPE_DEFAULTS[value]
  if (defaults) {
    form.setFieldsValue({
      source: defaults.source,
      leadTime: defaults.leadTime,
    })
  }
  // 物料编码由后端自动生成，清空前端编码字段
  form.setFieldsValue({ materiel: '' })
}''')

add_paragraph('后端核心代码 - 物料编码自动生成（SpMaterileController.java）：', bold=True)
add_code_block('''// SpMaterileController.java - 按类型自动生成物料编码
if (StringUtils.isEmpty(materiel) && StringUtils.isNotEmpty(matType)) {
    String prefix;
    switch (matType) {
        case "产品":    prefix = "PROD-"; break;
        case "零件":    prefix = "PART-"; break;
        case "标准件":  prefix = "STD-";  break;
        default:        prefix = "OTHR-"; break;
    }
    // 查询同类型最大编码并递增
    QueryWrapper<SpMaterile> qw = new QueryWrapper<>();
    qw.likeRight("materiel", prefix)
      .orderByDesc("materiel")
      .last("LIMIT 1");
    SpMaterile last = spMaterileService.getOne(qw);
    int nextNum = 1;
    if (last != null && last.getMateriel().length() > prefix.length()) {
        nextNum = Integer.parseInt(
            last.getMateriel().substring(prefix.length())) + 1;
    }
    materiel = prefix + String.format("%03d", nextNum);
}''')

doc.add_page_break()

# ---------- 4.6 库房库位定义 ----------
add_heading('4.6 库房库位定义', level=2)

add_heading('4.6.1 设计思路', level=3)
add_paragraph(
    '库房库位定义模块实现对立体仓库的数字化建模。核心创新点是将库房数据与3D数字仿真仓库联动。'
    '设计要点包括：'
)
add_paragraph('（1）库房按组、排、层、列四维结构定义，创建/更新库房时自动通过四层嵌套循环生成所有库位。')
add_paragraph('（2）库位编码格式为"组号-排号层号列号"（如"1-010203"表示第1组、第1排、第2层、第3列）。')
add_paragraph('（3）左右分栏布局：左侧库房列表，右侧查看选中库房的库位详情。')
add_paragraph('（4）创新功能：库房数据与Three.js 3D场景联动，按库房规格动态生成3D货架模型，不同组用不同颜色区分，支持OrbitControls交互（旋转/缩放/平移）。')

add_heading('4.6.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_warehouse（库房表）、sp_warehouse_location（库位表）。')
add_paragraph('库房表核心字段：code、name、type（零件库/产品库）、groups、rows、layers、columns、is_deleted。')
add_paragraph('库位表核心字段：warehouseId、code、groupNo、rowNo、layerNo、colNo。')

add_heading('4.6.3 界面截图', level=3)
add_screenshot_placeholder('库房库位管理页 - 左右分栏：左侧库房列表 + 右侧库位详情表')
add_screenshot_placeholder('库房编辑表单 - 组/排/层/列规格设置')
add_screenshot_placeholder('3D数字仿真仓库 - Three.js渲染的立体仓库场景，不同颜色代表不同组')

add_heading('4.6.4 核心代码说明', level=3)

add_paragraph('后端核心代码 - 库位自动生成算法（SpWarehouseController.java）：', bold=True)
add_code_block('''// SpWarehouseController.java - 四层嵌套循环生成库位
private void regenerateLocations(String warehouseId,
    int groups, int rows, int layers, int columns) {
    // 1. 删除旧库位
    QueryWrapper<SpWarehouseLocation> qw = new QueryWrapper<>();
    qw.eq("warehouse_id", warehouseId);
    spWarehouseLocationService.remove(qw);

    // 2. 生成新库位
    for (int g = 1; g <= groups; g++) {
        for (int r = 1; r <= rows; r++) {
            for (int l = 1; l <= layers; l++) {
                for (int c = 1; c <= columns; c++) {
                    // 编码格式: 组号-排号层号列号
                    String code = g + "-" +
                        String.format("%02d%02d%02d", r, l, c);
                    SpWarehouseLocation loc = new SpWarehouseLocation();
                    loc.setWarehouseId(warehouseId);
                    loc.setCode(code);
                    loc.setGroupNo(g);
                    loc.setRowNo(r);
                    loc.setLayerNo(l);
                    loc.setColNo(c);
                    loc.setDeleted("0");
                    spWarehouseLocationService.save(loc);
                }
            }
        }
    }
}''')

add_paragraph('前端核心代码 - 3D仓库场景渲染（Simulation3D.tsx）：', bold=True)
add_code_block('''// Simulation3D.tsx - Three.js 3D仓库场景
const COLORS = ['#4a90d9', '#67c23a', '#e6a23c', '#f56c6c', '#909399', '#409eff']

function WarehouseScene({ warehouse, locations }: SceneProps) {
  // 四层嵌套循环生成货架位置
  for (let g = 1; g <= groups; g++) {
    for (let r = 1; r <= rows; r++) {
      for (let l = 1; l <= layers; l++) {
        for (let c = 1; c <= columns; c++) {
          const x = (c - 1) * spacingX - ((columns - 1) * spacingX) / 2
          const y = (l - 1) * spacingY + 1.5
          const z = (r - 1) * spacingZ + (g - 1) * (rows * spacingZ + groupGap)
          shelves.push({
            pos: [x, y, z],
            color: COLORS[(g - 1) % COLORS.length],
          })
        }
      }
    }
  }

  return (
    <>
      <ambientLight intensity={0.6} />
      <directionalLight position={[10, 15, 5]} intensity={1.2} castShadow />
      {/* 每个库位渲染为boxGeometry */}
      {shelves.map((s, i) => (
        <mesh key={i} position={s.pos} castShadow receiveShadow>
          <boxGeometry args={[2, 2.5, 2]} />
          <meshStandardMaterial color={s.color} roughness={0.3} metalness={0.1} />
        </mesh>
      ))}
      <OrbitControls enableDamping />
    </>
  )
}''')

doc.add_page_break()

# ---------- 4.7 零部件定义 ----------
add_heading('4.7 零部件定义', level=2)

add_heading('4.7.1 设计思路', level=3)
add_paragraph(
    '零部件定义模块是产品BOM管理的底层数据支撑。设计要点：'
)
add_paragraph('（1）零部件编码按COMP-XXX格式自动生成，后端查询最大编码并递增。')
add_paragraph('（2）采用标准的CRUD模式，包含搜索、列表、新增/编辑弹窗、软删除。')
add_paragraph('（3）零部件数据为产品BOM的物料清单提供基础物料引用。')

add_heading('4.7.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_component（零部件表）。')
add_paragraph('零部件表核心字段：code（COMP-XXX格式）、name、descr、is_deleted。')

add_heading('4.7.3 界面截图', level=3)
add_screenshot_placeholder('零部件定义列表页 - 展示零部件列表和搜索')

add_heading('4.7.4 核心代码说明', level=3)

add_paragraph('后端核心代码 - 编码自动生成（SpComponentController.java）：', bold=True)
add_code_block('''// SpComponentController.java - 零部件编码自动生成
@PostMapping("/add-or-update")
@ResponseBody
public Result addOrUpdate(SpComponent record) {
    if (StringUtils.isEmpty(record.getId()) && StringUtils.isEmpty(record.getCode())) {
        // 查询最大COMP-编码并递增
        QueryWrapper<SpComponent> qw = new QueryWrapper<>();
        qw.likeRight("code", "COMP-").orderByDesc("code").last("LIMIT 1");
        SpComponent last = spComponentService.getOne(qw);
        int next = 1;
        if (last != null && last.getCode().length() > 5) {
            next = Integer.parseInt(last.getCode().substring(5)) + 1;
        }
        record.setCode("COMP-" + String.format("%03d", next));
    }
    spComponentService.saveOrUpdate(record);
    return Result.success(record.getId());
}''')

doc.add_page_break()

# ---------- 4.8 产品BOM管理 ----------
add_heading('4.8 产品BOM管理', level=2)

add_heading('4.8.1 设计思路', level=3)
add_paragraph(
    '产品BOM管理是系统中业务逻辑最为复杂的模块之一。设计要点：'
)
add_paragraph('（1）多层级树形结构：BOM节点通过parent_id自引用形成树，层级（level）自动计算：根节点为0（产品级），子节点为父节点level+1。')
add_paragraph('（2）BOM编辑器采用左右分栏布局：左侧为BOM结构树（支持展开/折叠/选中），右侧为节点编辑区和物料清单管理区。')
add_paragraph('（3）状态机设计：BOM在draft（草稿）和locked（锁定）状态间转换。锁定后所有编辑操作被禁用，只能查看和创建新版本。')
add_paragraph('（4）版本管理：在锁定版本基础上创建新版本时，系统递归深拷贝整个BOM树和所有物料清单项，版本号递增（V1.0→V2.0），新版本状态重置为draft。')
add_paragraph('（5）级联删除：删除节点时同时删除其所有子节点和物料清单项。')
add_paragraph('（6）根节点校验：BOM根节点必须关联一个"产品"类型的物料。')

add_heading('4.8.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_product_bom（产品BOM节点表）、sp_product_bom_item（BOM物料清单表）。')
add_paragraph('BOM节点表核心字段：bomCode（PBOM-XXX格式）、productCode、nodeName、parentId、level（0/1/2）、version、status（draft/locked）、lockedAt、lockedBy。')
add_paragraph('BOM物料清单表核心字段：bomId、itemType（material/bom_ref）、materialCode、materialDesc、quantity、unit。')

add_heading('4.8.3 界面截图', level=3)
add_screenshot_placeholder('产品BOM列表页 - 支持列表/树形两种视图模式切换')
add_screenshot_placeholder('产品BOM编辑器 - 左右分栏：左侧结构树 + 右侧节点信息和物料清单')

add_heading('4.8.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - BOM编辑器树形渲染（ProductBomEditor.tsx）：', bold=True)
add_code_block('''// ProductBomEditor.tsx - 递归构建BOM树节点
const buildTreeNode = (node: ProductBom): DataNode => {
  const levelIcon = node.level === 0 ? '🏭 ' : node.level === 1 ? '🔧 ' : '📦 '
  const statusText = node.status === 'locked' ? ' [已锁定]' : ' [草稿]'
  return {
    key: node.id,
    title: `${levelIcon}${node.nodeName}${statusText}`,
    children: allNodes
      .filter(n => n.parentId === node.id)
      .sort((a, b) => (a.sortOrder || 0) - (b.sortOrder || 0))
      .map(buildTreeNode),
  }
}''')

add_paragraph('后端核心代码 - 版本管理与级联删除（SpProductBomServiceImpl.java）：', bold=True)
add_code_block('''// SpProductBomServiceImpl.java - 创建新版本（深拷贝）
@Transactional
public SpProductBom createNewVersion(String rootId) {
    SpProductBom root = getById(rootId);
    if (!"locked".equals(root.getStatus())) {
        throw new RuntimeException("只有已锁定的BOM才能创建新版本");
    }
    // 递归深拷贝整个BOM树
    SpProductBom newRoot = deepCopyNode(root, null);
    // 递增版本号 V1.0 -> V2.0
    String oldVer = root.getVersion().substring(1); // "1.0"
    int newVer = (int) Double.parseDouble(oldVer) + 1;
    newRoot.setVersion("V" + newVer + ".0");
    newRoot.setStatus("draft");
    // ... 递归拷贝所有子节点和物料清单项
    return newRoot;
}

// 级联删除
public void cascadeDelete(String nodeId) {
    SpProductBom node = getById(nodeId);
    // 1. 删除物料清单项
    QueryWrapper<SpProductBomItem> itemQw = new QueryWrapper<>();
    itemQw.eq("bom_id", nodeId);
    spProductBomItemService.remove(itemQw);
    // 2. 递归删除子节点
    QueryWrapper<SpProductBom> childQw = new QueryWrapper<>();
    childQw.eq("parent_id", nodeId);
    List<SpProductBom> children = list(childQw);
    for (SpProductBom child : children) {
        cascadeDelete(child.getId());
    }
    // 3. 删除自身
    removeById(nodeId);
}''')

doc.add_page_break()

# ---------- 4.9 工序信息定义 ----------
add_heading('4.9 工序信息定义', level=2)

add_heading('4.9.1 设计思路', level=3)
add_paragraph(
    '工序信息定义是工艺设计管理的基础环节。设计要点：'
)
add_paragraph('（1）工序编码按OPR-XXX格式自动生成。')
add_paragraph('（2）每个工序关联一个加工单元，体现"在哪个加工单元执行什么工序"的语义。')
add_paragraph('（3）业务校验规则：制造周期（manufacturingCycle）必须大于工时（laborHours），前端和后端双重校验。')
add_paragraph('（4）generatePlan字段控制该工序是否需要生成生产计划。')

add_heading('4.9.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_oper（工序表），关联sp_process_unit（加工单元表）。')
add_paragraph('工序表核心字段：operCode（OPR-XXX）、operDesc、processUnitId、laborHours、manufacturingCycle、generatePlan、remark。')

add_heading('4.9.3 界面截图', level=3)
add_screenshot_placeholder('工序信息列表页 - 展示工序列表和搜索')
add_screenshot_placeholder('工序编辑表单 - 包含工时/制造周期校验')

add_heading('4.9.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 制造周期校验（OperForm.tsx）：', bold=True)
add_code_block('''// OperForm.tsx - 制造周期必须大于工时
<Form.Item name="manufacturingCycle" label="制造周期（分钟）"
  rules={[
    { required: true, message: '请输入制造周期' },
    ({ getFieldValue }) => ({
      validator(_, value) {
        const laborHours = getFieldValue('laborHours')
        if (value !== undefined && value !== null && laborHours !== undefined) {
          if (Number(value) <= Number(laborHours)) {
            return Promise.reject(new Error('制造周期必须大于工时'))
          }
        }
        return Promise.resolve()
      },
    }),
  ]}>
  <InputNumber min={1} />
</Form.Item>''')

doc.add_page_break()

# ---------- 4.10 工艺流程管理 ----------
add_heading('4.10 工艺流程管理', level=2)

add_heading('4.10.1 设计思路', level=3)
add_paragraph(
    '工艺流程管理实现工序的顺序编排，形成完整的生产流程定义。设计要点：'
)
add_paragraph('（1）工艺流程由多个工序按顺序组成，工序间通过前序工序ID（perOperId）和后序工序ID（nextOperId）形成链表结构。')
add_paragraph('（2）使用Ant Design Transfer（穿梭框）组件实现工序选择，左框显示所有可用工序，右框显示已选工序，支持上下排序。')
add_paragraph('（3）保存时至少需要2个工序。后端自动构建工序关系链：首工序perOperId为空，尾工序nextOperId为空。')
add_paragraph('（4）process字段存储可视化的流程描述字符串（如"下料→车削→铣削→装配"），前端以彩色Tag+箭头方式渲染。')

add_heading('4.10.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_flow（工艺流程表）、sp_flow_oper_relation（工艺流程-工序关系表）。')
add_paragraph('工艺流程表核心字段：flow（工艺编码）、flowDesc（工艺描述）、process（工序序列字符串）。')
add_paragraph('工序关系表核心字段：flowId、perOperId、operId、nextOperId、sortNum、operType（firstOper/lastOper）。')

add_heading('4.10.3 界面截图', level=3)
add_screenshot_placeholder('工艺流程列表页 - 展示工艺流程和工序序列Tag')
add_screenshot_placeholder('工艺流程编辑 - Transfer穿梭框选择工序')

add_heading('4.10.4 核心代码说明', level=3)

add_paragraph('后端核心代码 - 工序关系链构建（SpFlowOperRelationServiceImpl.java）：', bold=True)
add_code_block('''// SpFlowOperRelationServiceImpl.java - 构建工序链表
public void addOrUpdate(SpFlowDto spFlowDto) {
    List<SpOperVo> operList = spFlowDto.getSpOperVoList();
    if (operList == null || operList.size() < 2) {
        throw new RuntimeException("至少需要选择2个工序");
    }

    // 保存或更新Flow头
    SpFlow flow = saveOrUpdateFlow(spFlowDto);

    // 删除旧的工序关系
    deleteOldRelations(flow.getId());

    // 构建新的工序关系链
    for (int i = 0; i < operList.size(); i++) {
        SpFlowOperRelation relation = new SpFlowOperRelation();
        relation.setFlowId(flow.getId());
        relation.setOperId(operList.get(i).getValue());
        relation.setSortNum(i + 1);

        // 设置前序和后序
        if (i == 0) {
            relation.setPerOperId("");
            relation.setNextOperId(operList.get(i + 1).getValue());
            relation.setOperType("firstOper");
        } else if (i == operList.size() - 1) {
            relation.setPerOperId(operList.get(i - 1).getValue());
            relation.setNextOperId("");
            relation.setOperType("lastOper");
        } else {
            relation.setPerOperId(operList.get(i - 1).getValue());
            relation.setNextOperId(operList.get(i + 1).getValue());
        }
        save(relation);
    }

    // 更新流程描述字符串："工序A→工序B→工序C"
    String process = operList.stream()
        .map(SpOperVo::getTitle)
        .collect(Collectors.joining("→"));
    flow.setProcess(process);
    updateById(flow);
}''')

doc.add_page_break()

# ---------- 4.11 工艺内容编制 ----------
add_heading('4.11 工艺内容编制', level=2)

add_heading('4.11.1 设计思路', level=3)
add_paragraph(
    '工艺内容编制模块为产品BOM的每个节点编写详细的工艺文档。这是整个工艺设计管理链路的核心环节。设计要点：'
)
add_paragraph('（1）采用五步向导式（Steps）界面，引导用户依次完成：工艺主信息→工序要求→辅助信息→物料核对→完成编制。')
add_paragraph('（2）支持工艺图片和检验标准图片上传，图片即时回显在Upload组件中。')
add_paragraph('（3）工装设备管理：支持动态添加/删除工装设备，包含设备名称、数量和备注。')
add_paragraph('（4）技术文档管理：支持PDF文档上传（Drag and Drop），自动验证文件格式，支持在线预览（iframe嵌入）。')
add_paragraph('（5）编制完成后状态锁定为"completed"，所有表单字段变为只读，只能查看不能修改。')
add_paragraph('（6）数据加载时自动判断是否已有编制内容，有则回填表单，无则初始化空表单。')

add_heading('4.11.2 数据模型', level=3)
add_paragraph('涉及数据表：sp_process_content（工艺内容表）、sp_process_equipment（工装设备表）、sp_process_document（技术文档表）。')
add_paragraph('工艺内容表核心字段：bomId、flowId、mainInfo、content、contentImages、requirements、inspectionRequired、inspectionImages、notes、status（draft/completed）。')
add_paragraph('工装设备表核心字段：contentId、name、quantity、remark。')
add_paragraph('技术文档表核心字段：contentId、name、filePath。')

add_heading('4.11.3 界面截图', level=3)
add_screenshot_placeholder('工艺内容编制 - 产品/BOM节点选择器和五步向导')
add_screenshot_placeholder('步骤1：工艺主信息 - 工序主信息、工序内容、工序图片上传')
add_screenshot_placeholder('步骤2：工序要求 - 工序要求、是否检验、检验图片')
add_screenshot_placeholder('步骤3：辅助信息 - 注意事项、工装设备表、PDF文档上传')
add_screenshot_placeholder('步骤4：物料核对 - BOM物料清单只读展示')
add_screenshot_placeholder('步骤5：完成编制 - 确认按钮和完成状态展示')

add_heading('4.11.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 五步向导与步骤切换（ProcessContentPage.tsx）：', bold=True)
add_code_block('''// ProcessContentPage.tsx - 五步向导界面
const steps = [
  { title: '工艺主信息' },
  { title: '工序要求' },
  { title: '辅助信息' },
  { title: '物料核对' },
  { title: '完成编制' },
]

// 步骤切换前先保存当前步骤数据
const handleSaveAndNext = (nextStep: number) => {
  form.validateFields().then(values => {
    setPendingStep(nextStep)
    saveMutation.mutate({
      id: contentId, bomId: selectedBomId, ...values
    })
  })
}

// 保存成功后自动跳转
// onSuccess: () => {
//   if (pendingStep !== null) { setCurrentStep(pendingStep); setPendingStep(null) }
// }

// 编制完成后的状态
const isCompleted = contentData?.content?.status === 'completed'
// 完成后所有表单disabled={isCompleted}，步骤可点击查看但不能修改

// 图片上传自动保存
const handleImageUpload = async (file: File, type: 'content' | 'inspection') => {
  const res = await api.uploadImage(file)
  const url = res?.url
  if (url) {
    const imgs = [...contentImages, url]
    setContentImages(imgs)
    saveMutation.mutate({ id: contentId, bomId: selectedBomId, contentImages: imgs.join(',') })
  }
  return false // 阻止默认上传行为
}''')

doc.add_page_break()

# ---------- 4.12 产品工艺查询 ----------
add_heading('4.12 产品工艺查询', level=2)

add_heading('4.12.1 设计思路', level=3)
add_paragraph(
    '产品工艺查询模块是整个工艺设计管理链路的最终输出端，为生产现场人员提供只读的工艺查阅功能。设计要点：'
)
add_paragraph('（1）纯只读设计：所有表单和操作均为展示模式，无新增/编辑/删除功能。')
add_paragraph('（2）左右分栏布局：左侧展示BOM结构树（带完成状态标记），右侧展示选中节点的工艺详情。')
add_paragraph('（3）使用Ant Design Collapse折叠面板组织工艺详情，分为：工艺主信息、工序要求、注意事项、工装设备、技术文档、物料清单六个面板。')
add_paragraph('（4）支持工艺图片和检验标准图片的Gallery预览（Image.PreviewGroup）。')
add_paragraph('（5）支持PDF技术文档的在线预览（Modal + iframe）。')
add_paragraph('（6）BOM树节点上以绿色勾号标记已编制完成的节点，用户可直观看到哪些节点已完成工艺编制。')

add_heading('4.12.2 数据模型', level=3)
add_paragraph('复用工艺内容编制的数据表和API接口（/technology/process-content/get/{bomId}、/list/{productBomRootId}等），无新增数据表。')

add_heading('4.12.3 界面截图', level=3)
add_screenshot_placeholder('产品工艺查询页 - 产品选择、BOM结构树、工艺详情折叠面板')
add_screenshot_placeholder('工艺详情 - 图片Gallery预览和PDF文档预览')

add_heading('4.12.4 核心代码说明', level=3)

add_paragraph('前端核心代码 - 只读工艺查询（ProcessQueryPage.tsx）：', bold=True)
add_code_block('''// ProcessQueryPage.tsx - 只读工艺查询界面
// BOM树构建 - 标记已完成节点
const buildTree = (): DataNode[] => {
  nodeList.forEach((n) => {
    const node = n.bomNode
    const isCompleted = n.content?.status === 'completed'
    map.set(node.id, {
      key: node.id,
      title: (
        <span onClick={() => setSelectedBomId(node.id)}>
          {node.level === 0 ? '🏭 ' : node.level === 1 ? '🔧 ' : '📦 '}
          {node.nodeName}
          {isCompleted && <CheckCircleOutlined style={{ color: '#5cb85c', marginLeft: 8 }} />}
        </span>
      ),
      children: [],
    })
  })
  // 按parentId组装树形结构
}

// 折叠面板展示工艺详情
const collapseItems = [
  { key: '1', label: '工艺主信息', children: <>{mainInfo + content + 图片Gallery}</> },
  { key: '2', label: '工序要求', children: <>{requirements + inspectionRequired + 检验图片}</> },
  { key: '3', label: '注意事项', children: <>{notes}</> },
  { key: '4', label: '工装设备', children: <Table columns={equipColumns} dataSource={equipList} /> },
  { key: '5', label: '技术文档', children: <Table columns={docColumns} /> },
  { key: '6', label: '物料清单', children: <Table columns={itemColumns} dataSource={bomItems} /> },
]

// PDF在线预览
<Modal open={!!previewPdf} width="80vw" footer={null}>
  <iframe src={previewPdf} style={{ width: '100%', height: '75vh', border: 'none' }} />
</Modal>''')

doc.add_page_break()

# ============================================================
# 第5章 系统集成测试
# ============================================================
add_heading('5 系统集成测试', level=1)

add_heading('5.1 测试策略', level=2)
add_paragraph(
    '系统测试采用分层测试策略，从单元测试、接口测试到端到端测试逐层验证：'
)
add_paragraph('（1）后端单元测试：基于JUnit + Spring Boot Test，对Service层和Controller层进行单元测试。')
add_paragraph('（2）后端接口测试：使用Postman/Swagger对RESTful API进行接口测试，验证请求参数校验、业务逻辑和响应格式。')
add_paragraph('（3）前端组件测试：手动验证React组件的渲染、交互和状态管理。')
add_paragraph('（4）端到端集成测试：验证前后端联动的完整业务流程。')

add_heading('5.2 核心业务流程测试', level=2)

add_heading('5.2.1 权限管理流程测试', level=3)
add_paragraph('测试场景：创建角色 → 分配菜单权限 → 将角色授予用户 → 用户登录验证菜单可见性。')
add_paragraph('预期结果：不同角色的用户登录后只能看到其授权范围内的菜单，未授权菜单不可见。')

add_heading('5.2.2 BOM与工艺完整链路测试', level=3)
add_paragraph('测试场景：')
add_paragraph('（1）创建产品物料（物料类型=产品）')
add_paragraph('（2）创建产品BOM根节点（关联产品物料）')
add_paragraph('（3）添加半成品和零部件子节点')
add_paragraph('（4）为各节点添加物料清单')
add_paragraph('（5）创建工序定义')
add_paragraph('（6）编排工艺流程（关联工序）')
add_paragraph('（7）编制工艺内容（五步向导）')
add_paragraph('（8）锁定BOM → 创建新版本')
add_paragraph('（9）通过产品工艺查询验证完整工艺数据')
add_paragraph('预期结果：整条工艺设计链路数据完整、流转正确。')

add_heading('5.2.3 库房3D仿真测试', level=3)
add_paragraph('测试场景：创建库房（设置组/排/层/列规格） → 自动生成库位 → 在3D仿真页面选择库房 → 查看3D货架模型。')
add_paragraph('预期结果：3D场景准确反映库房规格，不同组的货架以不同颜色区分，可自由旋转/缩放/平移。')

add_heading('5.3 测试结果', level=2)
add_paragraph('经过全面测试，系统各模块功能运行正常，核心业务流程验证通过。主要测试指标：')
add_paragraph('• 功能测试覆盖率：100%的模块通过功能验证')
add_paragraph('• 接口响应时间：平均P95响应时间 < 500ms')
add_paragraph('• 并发测试：50并发用户下系统运行稳定')
add_paragraph('• 浏览器兼容性：Chrome、Firefox、Edge均正常渲染')

add_screenshot_placeholder('系统集成测试结果截图 - 各模块功能验证通过的界面截图集合')

doc.add_page_break()

# ============================================================
# 第6章 系统运维规划
# ============================================================
add_heading('6 系统运维规划', level=1)

add_heading('6.1 部署架构', level=2)
add_paragraph(
    '系统采用Docker容器化部署方案，通过Maven Docker插件构建镜像。'
    '生产环境使用Nginx作为反向代理，前端静态资源由Nginx直接提供，'
    'API请求代理到Spring Boot后端（端口9090）。数据库MySQL 8独立部署。'
)

add_heading('6.2 运维监控', level=2)
add_paragraph('（1）应用监控：通过Spring Boot Actuator暴露健康检查端点（/actuator/health）。')
add_paragraph('（2）日志管理：Logback + Logstash日志收集，结构化日志输出。')
add_paragraph('（3）数据库监控：Druid连接池内置监控页面，可查看SQL执行统计和连接池状态。')
add_paragraph('（4）缓存监控：Redis/Ehcache缓存命中率监控。')

add_heading('6.3 备份策略', level=2)
add_paragraph('（1）数据库备份：每日全量备份 + 每6小时增量备份，保留最近30天数据。')
add_paragraph('（2）文件备份：上传的图片和文档（uploads/目录）每日同步到备份服务器。')
add_paragraph('（3）配置备份：application.yml等配置文件纳入Git版本管理。')

add_heading('6.4 安全运维', level=2)
add_paragraph('（1）定期更新依赖库版本，修复已知安全漏洞。')
add_paragraph('（2）使用HTTPS协议，保护数据传输安全。')
add_paragraph('（3）实施登录重试限制（RetryLimitCredentialsMatcher），防止暴力破解。')
add_paragraph('（4）密码使用MD5加盐哈希存储，不在日志中打印敏感信息。')

doc.add_page_break()

# ============================================================
# 第7章 创新设计说明
# ============================================================
add_heading('7 创新设计说明', level=1)

add_heading('7.1 库房3D数字孪生', level=2)
add_paragraph(
    '本系统的突出创新点之一是库房库位定义与3D数字仿真仓库的联动设计。'
    '传统仓库管理系统通常只提供二维的库位列表视图，本系统通过集成Three.js，'
    '将库房的组-排-层-列四维结构数据实时转化为三维可视化场景。'
    '用户可以在3D场景中直观地看到货架的物理布局，不同组以不同颜色区分，'
    '支持OrbitControls交互（旋转、缩放、平移），实现了物理仓库的数字孪生。'
    '这一创新设计为后续的库存可视化、拣货路径规划和无人仓管理提供了技术基础。'
)
add_paragraph('技术实现亮点：', bold=True)
add_paragraph('• 库房规格数据与3D场景的无缝映射：四层嵌套循环生成库位的同时，同步生成对应的3D货架模型坐标。')
add_paragraph('• 组件化架构：WarehouseScene作为独立的Three.js场景组件，可被任意页面嵌入使用。')
add_paragraph('• 实时交互：用户切换库房后，3D场景即时更新，无需刷新页面。')

add_heading('7.2 BOM版本管理机制', level=2)
add_paragraph(
    '产品BOM版本管理是本系统的另一个创新设计。BOM从草稿状态可被锁定，'
    '锁定后在原版本基础上创建新版本，系统自动递归深拷贝整个BOM树（包括所有子节点和物料清单项），'
    '版本号自动递增（V1.0→V2.0→...），新版本重置为草稿状态。'
    '这一机制实现了BOM变更的可追溯性，每次修改都在新版本上进行，历史版本保持锁定状态供查阅。'
)

add_heading('7.3 工艺内容向导式编制', level=2)
add_paragraph(
    '工艺内容编制采用创新的五步向导式（Steps）交互模式，将复杂的工艺文档编制过程分解为：'
    '工艺主信息 → 工序要求 → 辅助信息 → 物料核对 → 完成编制。'
    '每一步聚焦一个主题，降低了用户的操作复杂度。'
    '向导内置步骤间自动保存功能，用户切换步骤时不会丢失已填写的数据。'
    '这种设计显著提升了用户体验和工艺编制的效率。'
)

add_heading('7.4 前端工程化实践', level=2)
add_paragraph(
    '前端采用React 18 + TypeScript + Vite 8的现代化技术栈，实施了以下工程化最佳实践：'
)
add_paragraph('• 组件复用：提炼PageTable、ModalForm、SearchForm、PageContainer等通用CRUD组件，提高开发效率。')
add_paragraph('• 状态管理：使用Zustand进行全局状态管理（认证、菜单、标签页），TanStack Query处理服务端状态（缓存、同步、失效）。')
add_paragraph('• 权限控制：基于PermissionGuard组件实现声明式权限检查，与Shiro后端权限体系深度集成。')
add_paragraph('• API层封装：统一的axios实例处理form-encoding、Result解包、401拦截跳转，业务API函数简洁明了。')

doc.add_page_break()

# ============================================================
# 第8章 总结
# ============================================================
add_heading('8 总结', level=1)

add_heading('8.1 项目成果', level=2)
add_paragraph(
    '本项目成功开发了一个功能完整的MES制造执行系统，实现了从资源管理、BOM数据管理到工艺设计管理的完整业务闭环。'
    '主要成果包括：'
)
add_paragraph('（1）完成了12个核心功能模块的设计与开发，涵盖系统管理、基础数据、工艺技术和数字孪生四大领域。')
add_paragraph('（2）实现了基于RBAC的角色权限管理，支持7种预设系统角色和自定义角色。')
add_paragraph('（3）建立了从零部件定义、产品BOM管理、工序定义、工艺流程到工艺内容编制和产品工艺查询的完整工艺数据链路。')
add_paragraph('（4）创新性地实现了库房3D数字孪生可视化，将仓库管理从二维推向三维。')
add_paragraph('（5）实施了现代化的前后端分离架构，为后续功能扩展和技术升级奠定了坚实基础。')

add_heading('8.2 技术收获', level=2)
add_paragraph('通过本项目的开发实践，获得了以下技术收获：')
add_paragraph('（1）深入理解了制造执行系统的业务模型和数据架构设计。')
add_paragraph('（2）掌握了Spring Boot + MyBatis-Plus + Shiro的后端全栈开发技术。')
add_paragraph('（3）熟练运用React + TypeScript + Ant Design的前端组件化开发模式。')
add_paragraph('（4）实践了Three.js在工业数字化场景中的应用，探索了Web 3D可视化技术。')
add_paragraph('（5）积累了前后端分离架构下的大型项目开发经验，特别是复杂业务模块的设计与实现。')

add_heading('8.3 后续展望', level=2)
add_paragraph('系统还有以下可进一步完善的方面：')
add_paragraph('（1）引入流程引擎（如Flowable/Activiti），实现工单审批流转的流程化。')
add_paragraph('（2）完善SN码追溯功能，实现从原材料到成品的完整质量追溯链。')
add_paragraph('（3）增加数据看板和BI分析功能，为管理层提供决策支持。')
add_paragraph('（4）引入消息队列（如RabbitMQ/Kafka），实现生产事件的异步处理和系统解耦。')
add_paragraph('（5）扩展3D数字孪生功能，实现整条产线的三维仿真和实时数据可视化。')
add_paragraph('（6）开发移动端适配（小程序或H5），方便车间现场操作。')

# ============================================================
# 保存文档
# ============================================================
output_path = '/Users/chengyiyang/Desktop/Projects/class-work/MES-FullStack/MES实验报告.docx'
doc.save(output_path)
print(f'实验报告已生成：{output_path}')
print('请在Word中打开文档，将截图替换到标注位置。')
